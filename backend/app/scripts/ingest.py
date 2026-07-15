"""
Ingests DIGI_Sales_Intelligence_Knowledge_Pack.docx into Qdrant via LangChain.

No content is paraphrased or summarized — every chunk is built directly from the
verbatim table/paragraph text in the source document. Chunking is per-record
(one outlet, one SKU, one promotion, one policy rule, etc. per chunk) with a
prepended contextual header, so retrieval stays precise on a small, fact-dense
corpus. See conversation history for the chunking rationale.

Run from backend/: python -m app.scripts.ingest
"""

from pathlib import Path

import docx
from langchain_core.documents import Document
from langchain_qdrant import QdrantVectorStore, RetrievalMode

from app.core.config import settings
from app.services.vector_store import (
    DENSE_VECTOR_NAME,
    SPARSE_VECTOR_NAME,
    get_embeddings,
    get_sparse_embeddings,
)

DOC_PATH = (
    Path(__file__).resolve().parents[2]
    / "data"
    / "DIGI_Sales_Intelligence_Knowledge_Pack.docx"
)


def _row_text(headers: list[str], row) -> str:
    cells = [c.text.strip() for c in row.cells]
    return "\n".join(f"{h}: {v}" for h, v in zip(headers, cells) if v)


def _table_headers(table) -> list[str]:
    return [c.text.strip() for c in table.rows[0].cells]


def build_profile_doc(table, section: str) -> Document:
    """Small Field/Value tables (profile, midday/EOD snapshots) stay as one chunk."""
    headers = _table_headers(table)
    lines = [_row_text(headers, row) for row in table.rows[1:]]
    text = f"{section}\n" + "\n".join(lines)
    return Document(page_content=text, metadata={"section": section})


def build_record_docs(
    table, section: str, meta_key: str | None = None, meta_column: str | None = None
) -> list[Document]:
    """One chunk per data row — SKUs, promotions, outlets, route entries, etc."""
    headers = _table_headers(table)
    docs = []
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        primary = cells[1] if len(cells) > 1 else cells[0]
        text = f"{section} - {primary}\n" + _row_text(headers, row)
        metadata = {"section": section}
        if meta_key and meta_column and meta_column in headers:
            metadata[meta_key] = cells[headers.index(meta_column)]
        docs.append(Document(page_content=text, metadata=metadata))
    return docs


def build_policy_docs(paragraphs: list[str], section: str) -> list[Document]:
    return [
        Document(
            page_content=f"{section}\nPolicy rule {i + 1}: {rule}",
            metadata={"section": section, "rule_index": i + 1},
        )
        for i, rule in enumerate(paragraphs)
    ]


def load_documents() -> list[Document]:
    d = docx.Document(str(DOC_PATH))
    tables = d.tables

    docs: list[Document] = []

    docs.append(build_profile_doc(tables[0], "1. Salesperson & Territory Profile"))

    docs.extend(
        build_record_docs(
            tables[1], "2. Product & SKU Catalogue", "sku_code", "SKU Code"
        )
    )
    docs.extend(
        build_record_docs(
            tables[2], "3. Active Promotions & Schemes", "promo_code", "Code"
        )
    )
    docs.extend(
        build_record_docs(
            tables[3],
            "4. Outlet Master & Performance Snapshot",
            "outlet_id",
            "Outlet ID",
        )
    )
    docs.extend(
        build_record_docs(
            tables[4], "5. Key Outlet Details & Recommended Orders", "outlet_name", "Outlet"
        )
    )

    route_intro = (
        "Ahmed has 32 planned outlets today. The first 15 are priority outlets because "
        "their combined sales are down approximately 18% versus the recent four-week "
        "average. Seven reported availability issues. Priority outlets should preferably "
        "be visited before 1 PM."
    )
    docs.append(
        Document(
            page_content=f"6. Route & Visit Plan\n{route_intro}",
            metadata={"section": "6. Route & Visit Plan"},
        )
    )
    docs.extend(
        build_record_docs(tables[5], "6. Route & Visit Plan", "outlet_name", "Outlet")
    )

    docs.extend(
        build_record_docs(
            tables[6], "7. Competitor Intelligence Log", "outlet_name", "Outlet"
        )
    )

    policy_rules = [
        "Order Bookers may communicate and execute only approved promotions and schemes.",
        "Order Bookers must not commit additional discounts, free goods, credit terms, or rebates without authorized approval.",
        "Competitor offers should be recorded with competitor name, offer mechanic, outlet, date, and whether the offer appears general or selective.",
        "Repeated competitor activity across three or more outlets in a territory should be escalated to the Area Sales Manager.",
        "Recommended orders are advisory and should consider sales history, purchase frequency, seasonality, active promotions, and outlet potential.",
        "Route deviations should be recorded with a reason. Missed high-priority outlets should be moved to the next available priority slot.",
        "Customer or market feedback should be logged factually. Unverified claims must be marked as unconfirmed.",
    ]
    docs.extend(build_policy_docs(policy_rules, "8. Commercial Policy & Escalation Rules"))

    docs.append(build_profile_doc(tables[7], "9. Midday Performance Snapshot"))
    docs.append(build_profile_doc(tables[8], "10. End-of-Day Performance Snapshot"))

    docs.extend(
        build_record_docs(tables[9], "11. Sales Coaching & FAQ Knowledge", "faq_question", "Question")
    )

    return docs


def run_ingestion(force_recreate: bool = True) -> int:
    documents = load_documents()
    print(f"Prepared {len(documents)} chunks from {DOC_PATH.name}")

    QdrantVectorStore.from_documents(
        documents,
        embedding=get_embeddings(),
        sparse_embedding=get_sparse_embeddings(),
        url=settings.qdrant_url,
        collection_name=settings.qdrant_collection,
        retrieval_mode=RetrievalMode.HYBRID,
        vector_name=DENSE_VECTOR_NAME,
        sparse_vector_name=SPARSE_VECTOR_NAME,
        force_recreate=force_recreate,
    )
    print(f"Ingested into Qdrant collection '{settings.qdrant_collection}'")
    return len(documents)


if __name__ == "__main__":
    run_ingestion()
